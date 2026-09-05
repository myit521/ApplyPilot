"""DOCX 导出。

对应 docs/design.md 第 8.3、3.1 节：批准后的分区结构简历版本
按模板渲染为可编辑 DOCX，结构化 JSON 保留为版本事实源。
"""

from __future__ import annotations

from io import BytesIO

from docx import Document

from .schemas import ResumeSections


def render_docx(job_title: str, sections: ResumeSections) -> bytes:
    doc = Document()
    doc.add_heading(f"应聘简历 - {job_title}", level=0)

    if sections.education:
        doc.add_heading("教育背景", level=1)
        for claim in sections.education:
            doc.add_paragraph(claim.text, style="List Bullet")

    if sections.skills:
        doc.add_heading("专业技能", level=1)
        for claim in sections.skills:
            doc.add_paragraph(claim.text, style="List Bullet")

    if sections.experience:
        doc.add_heading("工作与项目经历", level=1)
        for claim in sections.experience:
            doc.add_paragraph(claim.text, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
